import os
from collections import defaultdict
import binascii
from hashlib import md5
from io import BytesIO, StringIO, TextIOWrapper
from typing import List, Dict, Tuple, BinaryIO

import docx
import shortuuid
from anytree import Node
from docx import Document
from django.conf import settings
from docx.styles.style import _ParagraphStyle, _CharacterStyle
from docx.text.paragraph import Paragraph
from docx.text.run import Run
from lxml import etree
from lxml.etree import tostring
from markdown import Markdown
from html import escape as html_escape

from home.consts import COLUMN_TEXT, COLUMN_START, COLUMN_END, COLUMN_SPEAKER, EU_COUNTRIES
from le_francais.settings.base import FILES_LE_FRANCAIS_HTTPS, \
    FILES_LE_FRANCAIS_SITENAME

def is_gpt_disabled(request):
    if request is None:
        return True
    session = request.session
    geoip_dict = session.get('geoip', None)
    if ((geoip_dict and (geoip_dict.get('country_code') in os.environ.get('BLOCKED_COUNTRY_CODES').split(',')
    or geoip_dict.get('country_code') is None)) or request.GET.get('test_gpt_disabled', None) == '1'):
        return True
    return False


def files_le_francais_url(path=''):
	if FILES_LE_FRANCAIS_HTTPS:
		return f'https://{FILES_LE_FRANCAIS_SITENAME}{path}'
	else:
		return f'http://{FILES_LE_FRANCAIS_SITENAME}{path}'


def get_signature(params, secret_key=settings.WALLET_ONE_SECRET_KEY):
    """
    Base64(Byte(MD5(Windows1251(Sort(Params) + SecretKey))))
    params - list of tuples [('WMI_CURRENCY_ID', 643), ('WMI_PAYMENT_AMOUNT', 10)]
    """
    icase_key = lambda s: str(s).lower()

    lists_by_keys = defaultdict(list)
    if len(params[0]) == 3:
        for key, value, id in params:
            lists_by_keys[key].append(value)
    else:
        for key, value in params:
            lists_by_keys[key].append(value)

    str_buff = b''
    for key in sorted(lists_by_keys, key=icase_key):
        for value in sorted(lists_by_keys[key], key=icase_key):
            str_buff += str(value).encode('1251')
    str_buff += secret_key.encode('1251')

    md5_string = md5(str_buff).digest()
    return (binascii.b2a_base64(md5_string)[:-1])

def message_left(n, form1='чашечка', form2='чашечки', form5='чашечек'):
    n10 = n%10
    n100 = n%100
    if n == 0:
        return 'Вы израсходовали последнюю чашечку :( Вы можете пополнить их запас в настройках профиля.'.format(form5)
    elif n10 == 1 and n100 != 11:
        return 'У вас есть ещё {0} {1}. Просмотреть свой остаток и пополнить запас вы можете в настройках профиля.'.format(str(n), form1)
    elif n10 in [2, 3, 4] and n100 not in [12, 13, 14]:
        return 'У вас есть ещё {0} {1}. Просмотреть свой остаток и пополнить запас вы можете в настройках профиля.'.format(str(n), form2)
    else:
        return 'У вас есть ещё {0} {1}. Просмотреть свой остаток и пополнить запас вы можете в настройках профиля.'.format(str(n), form5)

def message(n, form1=' чашечка', form2=' чашечки', form5=' чашечек'):
    n10 = n % 10
    n100 = n % 100
    if n10 == 1 and n100 != 11:
        return '{0} {1}'.format(str(n), form1)
    elif n10 in [2, 3, 4] and n100 not in [12, 13, 14]:
        return '{0} {1}'.format(str(n), form2)
    else:
        return '{0} {1}'.format(str(n), form5)


# Letter Avatar https://github.com/CraveFood/avinit

import re

from base64 import b64encode
from xml.sax.saxutils import escape as xml_escape

try:
    import cairosvg
except ImportError:
    cairosvg = None


DEFAULT_FONTS = [
    'HelveticaNeue-Light',
    'Helvetica Neue Light',
    'Helvetica Neue',
    'Helvetica',
    'Arial',
    'Lucida Grande',
    'sans-serif',
]

DEFAULT_SETTINGS = {
    'width': '200',
    'height': '200',
    'radius': '0',
    'font-family': ','.join(DEFAULT_FONTS),
    'font-size': '80',
    'font-weight': '400',
}

SVG_TEMPLATE = """
<svg xmlns="http://www.w3.org/2000/svg" pointer-events="none"
     width="{width}" height="{height}">
  <rect width="{width}" height="{height}" style="{style}"></rect>
  <text text-anchor="middle" y="50%" x="50%" dy="0.35em"
        pointer-events="auto" fill="#ffffff" font-family="{font-family}"
        style="{text-style}">{text}</text>
</svg>
""".strip()
SVG_TEMPLATE = re.sub('(\s+|\n)', ' ', SVG_TEMPLATE)


DEFAULT_COLORS = [
    "#1abc9c", "#16a085", "#f1c40f", "#f39c12", "#2ecc71", "#27ae60",
    "#e67e22", "#d35400", "#3498db", "#2980b9", "#e74c3c", "#c0392b",
    "#9b59b6", "#8e44ad", "#bdc3c7", "#34495e", "#2c3e50", "#95a5a6",
    "#7f8c8d", "#ec87bf", "#d870ad", "#f69785", "#9ba37e", "#b49255",
    "#b49255", "#a94136",
]


def _from_dict_to_style(style_dict):
    return '; '.join(['{}: {}'.format(k, v) for k, v in style_dict.items()])


def _get_color(text, colors=None):
    if not colors:
        colors = DEFAULT_COLORS
    color_index = sum(map(ord, text)) % len(colors)
    return colors[color_index]


def get_svg_avatar(text, **kwargs):

    initials = '=)'

    text = text.strip()
    if text:
        split_text = text.split(' ')
        if len(split_text) > 1:
            initials = split_text[0][0] + split_text[-1][0]
        else:
            initials = split_text[0][0]

    opts = DEFAULT_SETTINGS.copy()
    opts.update(kwargs)

    style = {
        'fill': _get_color(text, opts.get('colors')),
        'width': opts.get('width') + 'px',
        'height': opts.get('height') + 'px',
        'rx': opts.get('radius') + 'px',
        '-moz-border-radius': opts.get('radius') + 'px',
    }

    text_style = {
        'font-weight': opts.get('font-weight'),
        'font-size': opts.get('font-size') + 'px',
    }

    return SVG_TEMPLATE.format(**{
        'height': opts.get('height'),
        'width': opts.get('width'),
        'style': _from_dict_to_style(style),
        'font-family': opts.get('font-family'),
        'text-style': _from_dict_to_style(text_style),
        'text': xml_escape(initials.upper()),
    }).replace('\n', '')


def get_png_avatar(text, output_file, **kwargs):
    if not cairosvg:
        raise Exception('CairoSVG is required to png conversions.')

    svg_avatar = get_svg_avatar(text, **kwargs)
    cairosvg.svg2png(svg_avatar, write_to=output_file)


def get_avatar_data_url(text, **kwargs):
    svg_avatar = get_svg_avatar(text, **kwargs)
    b64_avatar = b64encode(svg_avatar.encode('utf-8'))
    return 'data:image/svg+xml;base64,' + b64_avatar.decode('utf-8')

def get_navigation_object_from_page(root, current_page) -> dict:
    page_object = {
        "text": str(root.title),
        "nodes": [],
        "href": root.get_url(),
        "state": {}
    }
    from home.models import PageWithSidebar
    from home.models import LessonPage
    from home.models import ArticlePage
    from home.models import PodcastPage
    if (isinstance(root.specific, PageWithSidebar) or isinstance(root.specific, LessonPage) or isinstance(root.specific, ArticlePage) or isinstance(root.specific, PodcastPage)):
        menu_title = root.specific.menu_title
        if not isinstance(menu_title, str):
            menu_title = menu_title.decode()
        if menu_title != '':
            page_object["text"] = menu_title
        if not getattr(root.specific, 'is_selectable', True):
            page_object["selectable"] = False
    if root.id == current_page.id:
        page_object["state"] = {
            "selected": True
        }
        page_object["selectable"] = False
    for child in root.get_children():
        if child.show_in_menus and child.live:
            page_object["nodes"].append(
                get_navigation_object_from_page(child, current_page))
    if len(page_object["nodes"]) == 0:
        page_object.pop('nodes', None)
    return page_object

def get_nav_tree(root, current_page):
    if root.show_in_menus:
        nav_items = [get_navigation_object_from_page(root, current_page)]
    else:
        nav_items = get_navigation_object_from_page(root, current_page)["nodes"]
    return nav_items

def parse_tab_delimited_srt_file(file):
    from csv import DictReader, excel_tab
    from collections import OrderedDict
    result = OrderedDict()
    c = 0
    for row in DictReader(file, dialect=excel_tab):
        result[row[COLUMN_TEXT]] = {
            "id": f'line{c}',
            "start":row[COLUMN_START],
            "end":row[COLUMN_END],
            "speaker":row[COLUMN_SPEAKER]
        }
        c += 1
    #return OrderedDict(sorted(result.items(), key=lambda x: x[1]['start']))
    return result


def sub_html(html:str, parsed_srt):
    import re
    errors = []
    sub_map = []
    cursor=0
    for l, data in parsed_srt.items():
        new_line = f'<span class="transcript-line" ' \
                   f'id="{data["id"]}" ' \
                   f'data-start="{data["start"]}"' \
                   f' data-end="{data["end"]}"' \
                   f' data-speaker="{data["speaker"]}"' \
                   f'>{l}</span>'
        match = re.search(f'({re.escape(l)})', html[cursor:])
        if match is None:
            errors.append(f'NOT FOUND -- {l}')
            continue
        sub_map.append((match.start(1) + cursor, match.end(1) + cursor, new_line))
        cursor += match.end(1)
    for start, end, new_line in reversed(sub_map):
        html = html[:start] + new_line + html[end:]
    return html, errors

def create_document_from_transcript_srt(table_csv) -> BytesIO:
    parsed_table = parse_tab_delimited_srt_file(StringIO(table_csv))
    document = docx.Document()
    last_speaker = None
    for l, data in parsed_table.items():
        speaker = data['speaker']
        if speaker != last_speaker:
            current_p:Paragraph = document.add_paragraph('— ')
        else:
            current_p.add_run(' ')
        new_run:Run = current_p.add_run(f'{l}')
        comment = new_run.add_comment(
            text=f'{data["speaker"]} {data["start"]} {data["end"]} {data["id"]}'
        )
        last_speaker = speaker
    file = BytesIO()
    document.save(file)
    return file

def docx_parse_document(document):
    html = ""
    lines_map = []
    opened_tags = []
    toc_nodes = []
    is_list = False
    for p in document.paragraphs:
        html, lines_map, opened_tags, is_list = docx_parse_through_paragraph(document, p, html, lines_map, opened_tags, toc_nodes, is_list)
    html = re.sub('!\[(.+?)]\((.+?)\s"(.+?)"\)',
                  '<img title="\g<1>" src="\g<2>" alt="\g<3>">', html)
    html = re.sub('<p>\[HTML](.*?)\[/HTML]</p>', '\g<1>', html)
    return html, lines_map

def docx_parse_through_paragraph(document, paragraph:Paragraph, html, lines_map, opened_tags, toc_nodes, is_list=False):
    p_html = ''
    p_map = []
    runs_iterator = (run for run in paragraph.runs)

    opening_tags_html = ''
    if opened_tags:
        for tag in opened_tags:
            attrs = "".join((f' {attr}="{value}"' for attr, value in tag['attrs'].items()))
            opening_tags_html += f'<{tag["name"]}{attrs}>'

    p_html, p_map, opened_tags = xml_parse_through_element(document, runs_iterator, paragraph._element, p_html, p_map, opened_tags)

    closing_tags_html = ''
    if opened_tags:
        for tag in opened_tags:
            closing_tags_html += f'</{tag["name"]}>'

    # fixing out of place characters
    p_html = re.sub('([a-zA-Z])(<span.*?>)', '\g<2>\g<1>', p_html)
    p_html = re.sub('(</span>)([.,:])', '\g<2>\g<1>', p_html)

    p_style: _ParagraphStyle = paragraph.style

    if not 'List Paragraph' in p_style.name and is_list:
        html += '</ul>'
        is_list = False

    if 'Heading' in p_style.name or 'heading' in p_style.name:
        heading_value = int(re.match('[Hh]eading\s?(?P<value>\d)', p_style.name).group('value'))

        # TODO: TOC support
        # if heading_value > toc_nodes[-1].head_value:
        #     sub = Node(
        #         parent=toc_nodes[-1],
        #     )

        html += f'<h{heading_value + 1}>{opening_tags_html}{p_html}{closing_tags_html}</h{heading_value + 1}>'
    elif 'List Paragraph' in p_style.name:
        if not is_list:
            html += '<ul>'
            is_list = True
        html += f'<li>{opening_tags_html}{p_html}{closing_tags_html}</li>'
    else:
        html += f'<p>{opening_tags_html}{p_html}{closing_tags_html}</p>'
    lines_map += p_map
    return html, lines_map, opened_tags, is_list

def get_bookmark_name_id(e):
    s = tostring(e, encoding='utf-8').decode('utf-8')
    e_id = re.findall('w:id="(\d+)"', s)[0]
    e_name = re.findall("w:name=[\"'](.+?)[\"']", s)[0]
    return e_name, e_id

def xml_parse_through_element(document:Document, runs_iterator, element,
                              html:str, transcript_map:list, opened_tags, current_style=None):
    for sub in element.xpath('./*'):
        sub_name = sub.xpath('name()')
        if sub_name == 'w:t':
            t_html = html_escape(sub.xpath('text()')[0])
            if current_style['bold']:
                t_html = f'<strong>{t_html}</strong>'
            if current_style['italic']:
                t_html = f'<em>{t_html}</em>'
            if current_style['strikethrough']:
                t_html = f'<s>{t_html}</s>'
            if current_style['underline']:
                t_html = f'<u>{t_html}</u>'
            if current_style['is_superscript']:
                t_html = f'<sup>{t_html}</sup>'
            if current_style['is_subscript']:
                t_html = f'<sub>{t_html}<sub>'
            html += t_html
        elif sub_name == 'w:r':
            current_run:Run = next(runs_iterator)
            current_style = {
                'italic': current_run.italic,
                'bold': current_run.bold,
                'underline': current_run.underline,
                'strikethrough': current_run.font.strike,
                'is_hyperlink': current_run.is_hyperlink,
                'is_superscript': current_run.font.superscript,
                'is_subscript': current_run.font.subscript,
            }
            html, transcript_map, opened_tags = xml_parse_through_element(document, runs_iterator, sub, html,
                                             transcript_map, opened_tags, current_style)
        elif sub_name == 'w:bookmarkStart':
            bookmark_name, bookmark_id = get_bookmark_name_id(sub)
            if bookmark_name == '_GoBack':
                pass
            else:
                html += f'<span id="{bookmark_name}">'
            opened_tags.append({'name':'span', 'attrs':{'id': bookmark_name}})
        elif sub_name == 'w:bookmarkEnd':
            if opened_tags[-1]['attrs']['id'] == '_GoBack':
                pass
            else:
                html += f'</span>'
            if opened_tags[-1]['name'] == 'span':
                opened_tags.pop(-1)
            else:
                print(f'CAN NOT CLOSE TAG: SPAN')
        elif sub_name == 'w:commentRangeStart':
            comment_text = "".join(document.comments_part._element.get_comment_by_id(
                    int(sub.xpath('./@w:id')[0])).xpath('.//w:t/text()'))
            try:
                speaker, mm_start, mm_end = comment_text.rsplit(' ', 2)
                id_name = f"{speaker}{mm_start}{mm_end}"
                line_id = f'line-{shortuuid.uuid(name=id_name)}'
                transcript_map.append({
                    'start': mm_start, 'end': mm_end,
                    'id': line_id,
                    'speaker': speaker,
                })
                html += f'<span class="transcript-line" id="{line_id}" data-start="{mm_start}" data-end="{mm_end}" data-speaker="{speaker}">'
                opened_tags.append({'name':'span', 'attrs':{
                    'class':'transcript-line',
                    'id':line_id,
                    'data-start': mm_start,
                    'data-end': mm_end,
                    'data-speaker': speaker
                }})
            except ValueError:
                print('Wrong Comment')
                print(comment_text)
                opened_tags.append({'name': 'span', 'attrs': {'class': 'comment', 'data-text':comment_text}})
        elif sub_name == 'w:commentRangeEnd':
            html += '</span>'
            if opened_tags and opened_tags[-1]['name'] == 'span':
                opened_tags.pop(-1)
            else:
                print(f'CAN NOT CLOSE: SPAN')
        elif sub_name == 'w:instrText':
            # TODO parse instrText
            ...
    return html, transcript_map, opened_tags

def get_html_and_map_from_docx(docx_file):
    document = Document(docx_file)
    # paragraph: Paragraph
    # run: Run
    # start_ends_map = []
    # html = ""
    # c = 0
    # for paragraph in document.paragraphs:
    #     html += "<p>"
    #     # iterating through every xml element inside paragraph
    #     for elem in paragraph._element.xpath('./*'):
    #         c += 1
    #         if elem.xpath('name()') == 'w:commentRangeStart':
    #             try:
    #                 speaker, mm_start, mm_end = "".join(document.comments_part._element.get_comment_by_id(
    #                     int(elem.xpath('./@w:id')[0])).xpath('.//w:t/text()')).split(' ')[:3]
    #             except ValueError as e:
    #                 print("".join(document.comments_part._element.get_comment_by_id(
    #                     int(elem.xpath('./@w:id')[0])).xpath('.//w:t/text()')))
    #             start_ends_map.append({'start':mm_start, 'end':mm_end, 'id':f'line{c}'})
    #             html += f'<span class="transcript-line" id="line{c}" data-start="{mm_start}" data-end="{mm_end}">'
    #         elif elem.xpath('name()') == 'w:r':
    #             for r_elem in elem.xpath('./*'):
    #                 c += 1
    #                 if r_elem.xpath('name()') == 'w:commentRangeStart':
    #                     speaker, mm_start, mm_end = document.comments_part._element.get_comment_by_id(
    #                         int(r_elem.xpath('./@w:id')[0])).xpath('.//w:t/text()')[0].split(' ')[:3]
    #                     start_ends_map.append({'start': mm_start, 'end': mm_end, 'id': f'line{c}'})
    #                     html += f'<span class="transcript-line" id="line{c}" data-start="{mm_start}" data-end="{mm_end}">'
    #                 elif r_elem.xpath('name()') == 'w:commentRangeEnd':
    #                     html += f'</span>'
    #                 elif r_elem.xpath('name()') == 'w:instrText':
    #                     # TODO: INCLUDE PICTURE OR LINK OR WHATEVER
    #                     pass
    #                 else:
    #                     html += f'{"".join(r_elem.xpath(".//text()"))}'
    #         elif elem.xpath('name()') == 'w:commentRangeEnd':
    #             html += f'</span>'
    #     html += '</p>'
    # html = re.sub('!\[(.+?)\]\((.+?)\s"(.+?)"\)', '<img title="\g<1>" src="\g<2>" alt="\g<3>">', html)
    return docx_parse_document(document)

def parse_toc_file(file):
    headings = []
    for line in file.lines():
        text, value = line.split('\t')
        headings.append((text, value))
    return headings

def set_headings(html, headings:List[Tuple[str, int]]):
    for line, h_value in headings:
        html = html.replace(f'<p>{line}</p>', f'<h{h_value}>{line}</h>')
    return html

def parse_tags_file(file:TextIOWrapper):
    file.seek(0)
    styles = []
    for line in file.readlines():
        text, start_tag, end_tag = line.strip('\n').split('\t')
        styles.append((text, start_tag, end_tag))
    return styles

def set_tags(html, tags):
    for text, start_tag, end_tag in tags:
        html = html.replace(text, f'{start_tag}{text}{end_tag}')
    return html

def parse_docx(docx_file:BytesIO):
    document = Document(docx_file)
    starts_lines_map = {}
    start_ends_map = []
    html = ''
    inside_comment = True
    for paragraph in document.paragraphs:
        for elem in paragraph._element.xpath('./*'):
            if elem.xpath('name()') == 'w:commentRangeStart':
                speaker, mm_start, mm_end, l_id = document.comments_part._element.get_comment_by_id(
                    int(elem.xpath('./@w:id')[0])).xpath('.//w:t/text()')[0].split(' ')
                inside_comment = True
            elif elem.xpath('name()') == 'w:r':
                for r_elem in elem.xpath('./*'):
                    if r_elem.xpath('name()') == 'w:commentRangeStart':
                        speaker, mm_start, mm_end, l_id = document.comments_part._element.get_comment_by_id(
                            int(r_elem.xpath('./@w:id')[0])).xpath('.//w:t/text()')[0].split(' ')
                        inside_comment = True
                    elif r_elem.xpath('name()') == 'w:commentRangeEnd':
                        inside_comment = False
                    elif r_elem.xpath('name()') == 'w:instrText':
                        # TODO: INCLUDE PICTURE
                        pass
                    else:
                        if inside_comment:
                            ...
                        else:
                            ...
            elif elem.xpath('name()') == 'w:commentRangeEnd':
                inside_comment = False
        html += '</p>'
    html = re.sub('!\[(.+?)\]\((.+?)\s"(.+?)"\)', '<img title="\g<1>" src="\g<2>" alt="\g<3>">', html)
    return html, start_ends_map


def text_preprocessor(text:str):
    new_text = text.replace('--', '—')\
        .replace(' ?', ' ?')\
        .replace(' !', ' !')\
        .replace('- ', '— ')\
        .replace('...', '…').replace('..', '…')\
        .replace('[ ', '[').replace(' ]', ']')\
        .replace(' .', '.').replace(' ,', ',')
    # new_text = re.sub('—(\w)', '— \g<1>', new_text)
    # new_text = re.sub('(\w)—', '\g<1> —', new_text)
    new_text = re.sub('"([\w\d\-])', '«\g<1>', new_text)
    new_text = re.sub('([\w\d\-]|…|\.|!|\?)"', '\g<1>»', new_text)
    new_text = re.sub('(\S)\(', '\g<1> (', new_text)
    return new_text.lstrip(' ')


def parse_eaf(eaf_file:BinaryIO, sort_by_speaker) -> List[Dict]:
    eaf_file.seek(0)
    root = etree.parse(eaf_file)
    time_slots = {}
    for time_slot in root.xpath('./TIME_ORDER/TIME_SLOT'):
        time_slots[time_slot.xpath('@TIME_SLOT_ID')[0]] = int(time_slot.xpath('@TIME_VALUE')[0])
    transcript = []
    for tier in root.xpath('./TIER'):
        tier_id = str(tier.xpath('@TIER_ID')[0])
        for alignable_annotation in tier.xpath('.//ALIGNABLE_ANNOTATION'):
            time_start = time_slots[alignable_annotation.xpath('@TIME_SLOT_REF1')[0]]
            time_end = time_slots[alignable_annotation.xpath('@TIME_SLOT_REF2')[0]]
            texts = alignable_annotation.xpath('./ANNOTATION_VALUE/text()')
            if len(texts):
                text = "".join(texts)
            else:
                text = ""
            text = text_preprocessor(text)
            transcript.append({
                'speaker':tier_id,
                'start':time_start,
                'end':time_end,
                'text':text,
            })
    if sort_by_speaker:
        transcript.sort(key=lambda x: (x['speaker'], x['start']))
    else:
        transcript.sort(key=lambda x: x['start'])
    return transcript

def eaf_string_to_docx(s:str):
    if '\\\\' in s:
        r_pattern = re.compile(f"(?P<pre>.+)\\\\(?P<line>.+)\\\\(?P<post>.+)")
        match = re.search(r_pattern, s)
        pre_line, line, post_line = match.group('pre'), match.group('line'), match.group('post')
    else:
        pre_line, line, post_line = '', s, ''
    pre_line = pre_line.replace(r'\n', '\n')
    post_line = post_line.replace(r'\n', '\n')
    return pre_line, line, post_line

def map_to_docx(transcript_map, separate_speakers=True):
    document = docx.Document()
    last_speaker = None
    if not separate_speakers:
        current_p = document.add_paragraph('')
    for line in transcript_map:
        speaker, start, end, text = line['speaker'], line['start'], line['end'], line['text']
        if separate_speakers and speaker != last_speaker:
            current_p:Paragraph = document.add_paragraph('— ')
        else:
            pre_line, line, post_line = eaf_string_to_docx(text)
            current_p.add_run(' ') # FIXME referenced before assignment
        new_run = current_p.add_run(text)
        new_run.add_comment(f'{speaker} {start} {end}')
        last_speaker = speaker
    file = BytesIO()
    document.save(file)
    file.seek(0)
    return file


def eaf_to_docx(eaf_file:BinaryIO, separate_speakers=True, sort_by_speaker=False):
    transcript_map = parse_eaf(eaf_file, sort_by_speaker)
    return map_to_docx(transcript_map, separate_speakers)

def map_to_eaf(map:List[Dict], eaf_file:BytesIO):
    eaf_file.seek(0)
    root = etree.parse(eaf_file)
    map_tiers = list(set([l['speaker'] for l in map]))
    eaf_tiers = root.xpath('./TIER/@TIER_ID')
    # checking that all tiers from map are present in given eaf file
    if not set(map_tiers).issubset(set(eaf_tiers)):
        raise ValueError('Not all tiers are present in given eaf file')
    # creating time_slots to speaker annotations relating map
    eaf_annotations = []
    for annotation in root.xpath('./TIER/ANNOTATION/ALIGNABLE_ANNOTATION/.'):
        start_time_id = annotation.xpath('@TIME_SLOT_REF1')[0]
        stop_time_id = annotation.xpath('@TIME_SLOT_REF2')[0]
        eaf_annotations.append({
            'id': annotation.xpath('@ANNOTATION_ID')[0],
            'start':root.xpath(f'./TIME_ORDER/TIME_SLOT[@TIME_SLOT_ID={start_time_id}]/@TIME_VALUE')[0],
            'end':root.xpath(f'./TIME_ORDER/TIME_SLOT[@TIME_SLOT_ID={stop_time_id}]/@TIME_VALUE')[0],
            'line':annotation.xpath('./ANNOTATION_VALUE/text()')[0],
        })
    # deleting time order slots
    root.remove(root.xpath('../TIME_SLOT'))
    # deleting tiers
    for tier in map_tiers:
        ...
    # replace all tiers lines with new one
    # with кускуфештп related time order slots
    ...
    # inserting new time order slots
    ...
    # saving all changes to eaf file
    ...


def get_currency(request):
    if request.user.country_code in ['CA']:
        return 'cad'
    elif request.user.country_code in ['CH', 'LI']:
        return 'chf'
    elif request.user.countycode in ['UK']:
        return 'gbp'
    elif request.user.country_code in EU_COUNTRIES:
        return 'eur'
    elif request.user.country_code in ['US']:
        return 'usd'
    elif request.user.country_code in ['RU', 'BY']:
        return 'rub'
    else:
        return 'eur'
